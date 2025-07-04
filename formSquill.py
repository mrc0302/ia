import streamlit as st

# Configuração da página
# st.set_page_config(
#     page_title="Sistema de Modelos Judiciais",
#     page_icon="🧊", 
#     layout="wide",  
#     initial_sidebar_state="expanded"
# )

import time
from streamlit_quill import st_quill
import re
from database  import Database
import html
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import io
import os
from docx.enum.text import WD_COLOR_INDEX
import logging

# Configuração de logging
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger('FormSquill')

def main():

    def transform_results_to_items(results):
        # Agrupar por classe
        classes = {}
        for assunto, classe, preview in results:
            if classe not in classes:
                classes[classe] = []
            classes[classe].append((assunto, preview))
        
        # Criar itens para o accordion
        items = []
        for classe in sorted(classes.keys()):
            children = []
            for assunto, preview in sorted(classes[classe]):
                children.append({
                    "label": assunto,
                    "key": assunto,
                    "description": preview[:100] + "..." if preview else ""
                })
            
            items.append({
                "label": classe,
                "children": children,
                "key": f"class_{classe}"
            })
        
        return items

    def html_to_word(html_content):
        """Converts Quill HTML content to Word document preserving judicial document formatting"""
        try:
            doc = Document()
            if not html_content:
                return doc
                
            html_content = html.unescape(html_content)
            soup = BeautifulSoup(html_content, 'html.parser')

            # Define default paragraph style
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(12)
            
            def get_color_from_style(style_str):
                # Handle both RGB and Hex color formats
                color_match = re.search(r'color:\s*(#[0-9a-fA-F]{6}|rgb\(\d+,\s*\d+,\s*\d+\))', style_str)
                if color_match:
                    color = color_match.group(1)
                    if color.startswith('#'):
                        return tuple(int(color[i:i+2], 16) for i in (1, 3, 5))
                    return tuple(map(int, re.findall(r'\d+', color)))
                return None

            def process_text_with_styles(paragraph, element):
                if isinstance(element, str):
                    run = paragraph.add_run(element)
                    return

                # Get text content
                text = element.get_text()
                if not text.strip():
                    return

                # Create run with text
                run = paragraph.add_run(text)
                
                # Apply styling
                style_attr = element.get('style', '')
                
                # Handle text color
                rgb = get_color_from_style(style_attr)
                if rgb:
                    run.font.color.rgb = RGBColor(*rgb)
                
                # Handle background color (highlight)
                if 'background-color: yellow' in style_attr:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
                # Handle text formatting
                if element.name == 'strong' or 'ql-bold' in element.get('class', []):
                    run.bold = True
                if element.name == 'em' or 'ql-italic' in element.get('class', []):
                    run.italic = True
                if element.name == 'u' or 'ql-underline' in element.get('class', []):
                    run.underline = True
                if element.name == 's' or 'ql-strike' in element.get('class', []):
                    run.strike = True

            def process_paragraph(element):
                # Create new paragraph
                p = doc.add_paragraph()
                
                # Handle alignment
                if 'ql-align-center' in element.get('class', []):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif 'ql-align-right' in element.get('class', []):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif 'ql-align-justify' in element.get('class', []):
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Process content preserving formatting
                for child in element.children:
                    if isinstance(child, str):
                        run = p.add_run(child)
                    else:
                        process_text_with_styles(p, child)
                
                # Add spacing after paragraph
                p.spacing_after = Pt(12)
                return p

            # Process headers
            for header in soup.find_all(['h1', 'h2']):
                p = doc.add_paragraph()
                run = p.add_run(header.get_text())
                run.bold = True
                if header.name == 'h1':
                    run.font.size = Pt(16)
                else:
                    run.font.size = Pt(14)
                p.spacing_after = Pt(12)

            # Process all paragraphs and divs
            for element in soup.find_all(['p', 'div']):
                if element.get_text().strip():  # Skip empty paragraphs
                    process_paragraph(element)

            # Process lists
            for list_element in soup.find_all(['ul', 'ol']):
                list_style = 'bullet' if list_element.name == 'ul' else 'decimal'
                for item in list_element.find_all('li'):
                    p = doc.add_paragraph()
                    p.style = list_style
                    process_text_with_styles(p, item)

            # Add extra paragraph for spacing at end
            doc.add_paragraph()
            
            return doc
        except Exception as e:
            logger.error(f"Error creating Word document: {str(e)}")
            return None

    def generate_rtf_from_html(html_content):
        """
        Gera um conteúdo RTF básico a partir do HTML.
        Esta é uma conversão simples que preserva o texto base.
        """
        if not html_content:
            return ""
        
        # Remove tags HTML para obter texto simples
        soup = BeautifulSoup(html_content, 'html.parser')
        text = soup.get_text()
        
        # Cabeçalho RTF básico
        rtf_header = r"{\rtf1\ansi\ansicpg1252\deff0\deflang1046"
        rtf_footer = r"}"
        
        # Substitui quebras de linha por \par no RTF
        text = text.replace('\n', r'\par ')
        
        # Monta o RTF completo
        rtf_content = f"{rtf_header}\n{text}\n{rtf_footer}"
        
        return rtf_content

    def init_session_state():
        # Inicialização de todas as variáveis de sessão
        if "form_assunto" not in st.session_state:
            st.session_state.form_assunto = ""
        if "form_classe" not in st.session_state:
            st.session_state.form_classe = ""
        if "form_content" not in st.session_state:
            st.session_state.form_content = ""
        if "search_text" not in st.session_state:
            st.session_state.search_text = ""
        if "search_mode" not in st.session_state:
            st.session_state.search_mode = "Assunto"
        if "show_dialog" not in st.session_state:
            st.session_state.show_dialog = False
        if "action_type" not in st.session_state:
            st.session_state.action_type = None
        if "temp_data" not in st.session_state:
            st.session_state.temp_data = {}
        if "word_file" not in st.session_state:
            st.session_state.word_file = None
        if "show_word_download" not in st.session_state:
            st.session_state.show_word_download = False
        if "highlight_enabled" not in st.session_state:
            st.session_state.highlight_enabled = True
        if "search_results" not in st.session_state:
            st.session_state.search_results = None
        if 'dialog_open' not in st.session_state:
            st.session_state.dialog_open = False
        # Variáveis para mensagens de erro/sucesso
        if "success_message" not in st.session_state:
            st.session_state.success_message = None
        if "error_message" not in st.session_state:
            st.session_state.error_message = None

    def load_css(file_path):
        try:
            with open(file_path) as f:
                st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
        except Exception as e:
            logger.warning(f"Erro ao carregar CSS: {str(e)}")

    try:
        load_css("static/styles.css")
    except:
        logger.warning("Arquivo CSS não encontrado. Usando estilos padrão.")

    @st.dialog("Aviso", width="small") 
    def show_confirmation_export():      
            
        st.markdown(f"<br><div style='text-align: center; font-size: 20px;'>Documento exportado na pasta Downloads!</div><br>", unsafe_allow_html=True)
        action_display = {
            "ok": "ok"
        }.get(st.session_state.action_type, "")
                   
        col1 = st.columns(1)[0]
        
        # with col1:
        #     if st.button("Ok", key="confirm_yes"):
        #         if st.session_state.action_type == "ok":
        #             st.session_state.dialog_open = False
        #         st.rerun()

    @st.dialog("Aviso", width="small") 
    def show_error_dialog(error):      
        with st.container():
          
            st.markdown(f"<br><div style='text-align: center; font-size: 20px;'>{error}</div><br>", unsafe_allow_html=True)
            
            action_display = {
                "ok": "ok"
            }.get(st.session_state.action_type, "")
                
            
            # col1 = st.columns(1)[0]
            
            # with col1:
            #     if st.button("Ok", key="confirm_yes"):
            #         if st.session_state.action_type == "ok":
            #             st.session_state.dialog_open = False
            #         st.rerun()

    @st.dialog("Aviso",width="small") 
    def show_confirmation_dialog():
        #with st.container(key="dialog"):
            action_display = {
                "salvar": "salvar",
                "excluir": "excluir",
                "limpar": "limpar",
            }.get(st.session_state.action_type, "")
            
            
            st.markdown(f"<div style='text-align: center; font-size: 20px;'>Deseja realmente {action_display} este documento?</div>", unsafe_allow_html=True)
            
            col1_button, col2_button = st.columns([1, 1])
            
            with col1_button:
                
                if st.button("Sim", key="confirm_yes"):
                    if st.session_state.action_type == "salvar":
                        execute_save()
                    elif st.session_state.action_type == "excluir":
                        execute_delete()
                    elif st.session_state.action_type == "limpar":
                        execute_clear()
                    st.session_state.show_dialog = False
                    st.session_state.temp_data = {}
                    st.rerun()
            
            with col2_button:

                if st.button("Não", key="confirm_no"):
                    st.session_state.show_dialog = False
                    st.session_state.action_type = None
                    st.session_state.temp_data = {}
                    st.session_state.search_text = ""  # Limpa o resultado da busca
                    st.rerun()
    def execute_save():
        """Salva o documento no banco de dados"""
        try:
            app = Database()
            assunto = st.session_state.temp_data.get('assunto')
            classe = st.session_state.temp_data.get('classe')
            html_content = f"{st.session_state.temp_data.get('content')}".rstrip() 
            
            # Gera o conteúdo RTF a partir do HTML
            rtf_content = generate_rtf_from_html(html_content)
            
            content_to_save = {
                'html': html_content,
                'rtf': rtf_content
            }
            
            logger.info(f"Salvando documento: assunto={assunto}, classe={classe}")
            success, message = app.save_or_update_document(assunto, classe, content_to_save)
            
            if success:
                logger.info("Documento salvo com sucesso")
                set_form_data(assunto, classe, html_content)
                st.session_state.success_message = "Documento salvo com sucesso!"
            else:
                logger.error(f"Erro ao salvar documento: {message}")
                st.session_state.error_message = message
        
        except Exception as e:
            logger.error(f"Exceção ao salvar documento: {str(e)}")
            st.session_state.error_message = f"Erro ao salvar: {str(e)}"

    def execute_delete():
        """Exclui o documento do banco de dados"""
        try:
            app = Database()
            assunto = st.session_state.form_assunto

            if assunto:
                logger.info(f"Excluindo documento: assunto={assunto}")
                success, message = app.delete_document(assunto)
                
                if success:
                    logger.info("Documento excluído com sucesso")
                    set_form_data()
                    st.session_state.success_message = "Documento excluído com sucesso!"
                else:
                    logger.error(f"Erro ao excluir documento: {message}")
                    st.session_state.error_message = message
            else:
                logger.warning("Tentativa de excluir documento sem assunto")
                st.session_state.error_message = "Selecione um documento para excluir"
        
        except Exception as e:
            logger.error(f"Exceção ao excluir documento: {str(e)}")
            st.session_state.error_message = f"Erro ao excluir: {str(e)}"

    def execute_clear():
        """Limpa o formulário e campos relacionados"""
        logger.info("Limpando formulário")
        # Limpa os campos do formulário
        set_form_data()
        
        # Limpa o campo de busca e seus resultados
        st.session_state.search_text = ""
        st.session_state.search_mode = "Assunto"  # Reseta para o valor padrão
        st.session_state.search_results = None
        
        st.session_state.success_message = "Formulário limpo com sucesso!"

    def set_form_data(assunto="", classe="", content=""):
        """Define os dados do formulário na sessão"""
        st.session_state.form_assunto = assunto
        st.session_state.form_classe = classe
        blank_lines = "<BR>".join(["<BR>"] * 25)    
        st.session_state.form_content = content + blank_lines

    def highlight_search_terms(html_content, search_text):
        """Destaca termos de pesquisa no conteúdo HTML"""
        if not st.session_state.highlight_enabled or not search_text or not html_content:
            return html_content
        
        terms = [term.strip() for term in search_text.split(',')]
        terms = [term for term in terms if term]
        
        highlighted_content = html_content
        for term in terms:
            try:
                escaped_term = re.escape(term)
                pattern = re.compile(f'({escaped_term})', re.IGNORECASE)
                highlighted_content = pattern.sub(
                    lambda m: f'<span style="background-color: yellow;">{m.group(1)}</span>',
                    highlighted_content
                )
            except Exception as e:
                logger.warning(f"Erro ao destacar termo '{term}': {str(e)}")
        
        return highlighted_content

    def execute_search():
        """Executa a pesquisa no banco de dados"""
        app = Database()
        if st.session_state.search_text:
            try:
                logger.info(f"Pesquisando: modo={st.session_state.search_mode}, texto={st.session_state.search_text}")
                results, error_message = app.search_documents(
                    st.session_state.search_text, 
                    st.session_state.search_mode
                )
                
                if error_message:
                    logger.error(f"Erro na busca: {error_message}")
                    st.error(f"Erro na busca: {error_message}")
                    st.session_state.search_results = None
                elif not results:
                    logger.info("Nenhum documento encontrado")
                    st.warning("Nenhum documento encontrado!")
                    st.session_state.search_results = None
                else:
                    logger.info(f"Encontrados {len(results)} resultados")
                    st.success(f"Encontrados {len(results)} resultados")
                    st.session_state.search_results = results
                
            except Exception as e:
                logger.error(f"Exceção na busca: {str(e)}")
                st.error(f"Erro na busca: {str(e)}")
                st.session_state.search_results = None

    # Inicializa o estado da sessão
    init_session_state()
    app = Database()
    
    # Esconde o header da página
    st.markdown("""
        <style>
            .main {
                overflow: hidden;
            }
            .block-container {
                padding-top: 1rem;
                padding-bottom: 0rem;
                height: 100vh;
            }
            footer {
                display: none;
            }
            #MainMenu {
                visibility: hidden;
            }
            header {
                visibility: hidden;
            }                  
        </style>
    """, unsafe_allow_html=True)

    # # Exibir mensagens de sucesso ou erro
    # if st.session_state.success_message:
    #     st.success(st.session_state.success_message)
    #     st.session_state.success_message = None
        
    # if st.session_state.error_message:
    #     st.error(st.session_state.error_message)
    #     st.session_state.error_message = None

    # Mostrar diálogo de confirmação se necessário
    if st.session_state.show_dialog:
        show_confirmation_dialog()
            
    if st.session_state.dialog_open:
        st.success("Documento exportado na pasta Downloads!")
        st.session_state.dialog_open = False

    # Layout principal
    col1, col2 = st.columns([1.5, 5])

    with col1:
        st.markdown("<h4 style='text-align: left'> Pesquisa de Modelos</h4>", unsafe_allow_html=True)
        st.markdown("<br>", unsafe_allow_html=True)
        
        st.session_state.highlight_enabled = st.checkbox("Destacar termos encontrados", value=st.session_state.highlight_enabled)
        
        # Campo de pesquisa
        search_text = st.text_input("Parâmetro:", key="search_text")
        search_mode = st.radio("Modalidade:", 
                            ["Assunto", "Classe", "Todos", "Multi-termo"],
                            key="search_mode")
        
        # Botão de pesquisa
        if st.button("🔍 Pesquisar", type="primary", key="search"):
            execute_search()

        # Container para resultados da pesquisa
        container = st.container(height=500, key="container")    
    
        with container:
            if st.session_state.search_results is not None:
                classes = {}
                for assunto, classe, preview in st.session_state.search_results:
                    if classe not in classes:
                        classes[classe] = []
                    classes[classe].append((assunto, preview))

                for classe in sorted(classes.keys()):
                    with st.expander(f"► {classe}"):
                        for assunto, preview in sorted(classes[classe]):
                            if st.button(f"📝 {assunto}", type="secondary", key=f"btn_{classe}_{assunto}"):
                                doc, doc_message = app.get_document(assunto)
                                if doc:
                                    highlighted_content = highlight_search_terms(doc[3], search_text)
                                    set_form_data(doc[1], doc[2], highlighted_content)
                                else:
                                    st.error(doc_message)

    with col2:
        with st.form(key="form"):
            st.markdown("<h4 style='text-align: center; margin-top: -2rem; margin-bottom:0px'>📝Modelos de decisões judiciais</h4>", unsafe_allow_html=True)
            assunto = st.text_input("Assunto:", value=st.session_state.form_assunto)
            classe = st.text_input("Classe:", value=st.session_state.form_classe)
                       
            with st.container(height=600, border=True):
                toolbar = {
                    'container': [
                        [{'header': [1, 2, 3, 4, 5, 6, False]}],
                        ['bold', 'italic', 'underline', 'strike'],
                        ['blockquote', 'code-block'],
                        [{'list': 'ordered'}, {'list': 'bullet'}],
                        [{'script': 'sub'}, {'script': 'super'}],
                        [{'indent': '-1'}, {'indent': '+1'}],
                        [{'direction': 'rtl'}],
                        [{'size': ['small', False, 'large', 'huge']}],                        
                        [{'color': []}, {'background': []}],
                        [{'font': []}],
                        [{'align': []}],
                        ['clean'],
                        ['link', 'image', 'video'],
                        ['formula']
                    ]
                }
                
                if st.session_state.form_content == "":
                    blank_lines = "<BR>".join(["<BR>"] * 25)
                    st.session_state.form_content = blank_lines

                content = st_quill(
                    value=f"""<h2>{st.session_state.form_content}</h2>""",
                    html=True,
                    toolbar=toolbar,
                    history=True,
                    preserve_whitespace=True,
                    readonly=False,
                    key=None
                )
                    
            # Botões de ação
            col1, col2, col3, col4 = st.columns(4)

            with col1:
                if st.form_submit_button(" 💾 Salvar"):
                    if not assunto:
                        st.session_state.error_message = "O campo Assunto é obrigatório"
                    elif not classe:
                        st.session_state.error_message = "O campo Classe é obrigatório"
                    else:
                        st.session_state.show_dialog = True
                        st.session_state.action_type = "salvar"
                        st.session_state.temp_data = {
                            'assunto': assunto,
                            'classe': classe,
                            'content': content
                        }
                        st.rerun()

            with col2:
                if st.form_submit_button("✨ Novo"):
                    st.session_state.show_dialog = True
                    st.session_state.action_type = "limpar"             
                    st.rerun()

            with col3:
                if st.form_submit_button("🗑️ Excluir"):
                    if st.session_state.form_assunto:
                        st.session_state.show_dialog = True
                        st.session_state.action_type = "excluir"
                        st.rerun()
                    else:
                        st.session_state.error_message = "Selecione um documento para excluir"
                        st.rerun()

            with col4:
                if st.form_submit_button("📝 Exportar Word"):
                    if content:
                        try:
                            doc = html_to_word(content)
                            if doc:
                                docx_file = io.BytesIO()
                                doc.save(docx_file)
                                docx_file.seek(0)                            
                                filename = f"{assunto if assunto else 'documento'}.docx"
                                download_path = os.path.join(os.path.expanduser('~'), 'Downloads')
                                doc.save(os.path.join(download_path, filename))
                                st.session_state.dialog_open = True
                                st.rerun()
                        except Exception as e:
                            logger.error(f"Erro ao exportar para Word: {str(e)}")
                            st.session_state.error_message = f"Erro ao exportar: {str(e)}"
                            st.rerun()

if __name__ == "__main__":
    main()
