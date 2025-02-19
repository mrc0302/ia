import streamlit as st
import google.generativeai as genai
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import os
import tempfile
from dotenv import load_dotenv
import pandas as pd
from langchain_core.documents import Document
from typing import List
import docx
import PyPDF2
from bs4 import BeautifulSoup
import io
import shutil
import logging
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
import pickle

# Configuração de página como primeiro comando
st.set_page_config(
    page_title="RAG App com Google Drive",
    layout="wide"
)

# Configuração de logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
logger.info("Iniciando aplicação")

# Inicializações globais
if 'vector_store' not in st.session_state:
    st.session_state['vector_store'] = None
if 'temp_dir' not in st.session_state:
    st.session_state['temp_dir'] = tempfile.mkdtemp()
if 'drive_service' not in st.session_state:
    st.session_state['drive_service'] = None

# Funções para Google Drive
def authorize_drive():
    """Autoriza acesso ao Google Drive"""
    SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
    creds = None
    
    # Verifica se há token salvo
    token_upload = st.file_uploader("Faça upload do arquivo token.pickle (se já tiver)", type=['pickle'])
    if token_upload:
        token_bytes = token_upload.getvalue()
        creds = pickle.loads(token_bytes)
    
    # Se não há credenciais válidas, pede autenticação
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            # Precisa fazer upload do credentials.json
            creds_upload = st.file_uploader("Faça upload do arquivo credentials.json do Google Cloud", type=['json'])
            if not creds_upload:
                st.warning("Você precisa fazer upload do arquivo credentials.json para acessar o Google Drive")
                return None
                
            # Salva credentials.json temporariamente
            creds_path = os.path.join(st.session_state['temp_dir'], 'credentials.json')
            with open(creds_path, 'wb') as f:
                f.write(creds_upload.getvalue())
                
            try:
                flow = InstalledAppFlow.from_client_secrets_file(creds_path, SCOPES)
                auth_url = flow.authorization_url()[0]
                st.markdown(f"[Clique aqui para autorizar]({auth_url})")
                code = st.text_input("Cole o código de autorização:")
                if code:
                    flow.fetch_token(code=code)
                    creds = flow.credentials
                    # Salva o token para uso futuro
                    token_path = os.path.join(st.session_state['temp_dir'], 'token.pickle')
                    with open(token_path, 'wb') as token:
                        pickle.dump(creds, token)
                    # Disponibiliza o download do token
                    with open(token_path, 'rb') as token_file:
                        token_bytes = token_file.read()
                        st.download_button(
                            "Baixar token para uso futuro",
                            token_bytes,
                            "token.pickle",
                            "application/octet-stream"
                        )
                else:
                    return None
            except Exception as e:
                st.error(f"Erro na autenticação: {str(e)}")
                return None
                
    # Constrói o serviço Drive
    try:
        service = build('drive', 'v3', credentials=creds)
        return service
    except Exception as e:
        st.error(f"Erro ao construir serviço Drive: {str(e)}")
        return None

def list_drive_folders(service, parent_id='root'):
    """Lista pastas no Google Drive"""
    try:
        query = f"mimeType='application/vnd.google-apps.folder' and '{parent_id}' in parents and trashed=false"
        results = service.files().list(
            q=query,
            spaces='drive',
            fields='files(id, name)'
        ).execute()
        folders = results.get('files', [])
        return folders
    except Exception as e:
        st.error(f"Erro ao listar pastas: {str(e)}")
        return []

def list_drive_folders_recursive(service, parent_id='root', prefix=''):
    """Lista pastas recursivamente com formato hierárquico"""
    folders_flat = {}
    try:
        folders = list_drive_folders(service, parent_id)
        for folder in folders:
            folder_name = f"{prefix}{folder['name']}"
            folders_flat[folder_name] = folder['id']
            # Busca subpastas
            subfolders = list_drive_folders_recursive(
                service, 
                folder['id'], 
                prefix=f"{folder_name}/"
            )
            folders_flat.update(subfolders)
        return folders_flat
    except Exception as e:
        st.error(f"Erro ao listar pastas recursivamente: {str(e)}")
        return {}

def download_file_from_drive(service, file_id, destination):
    """Baixa arquivo do Google Drive"""
    try:
        request = service.files().get_media(fileId=file_id)
        with open(destination, 'wb') as f:
            downloader = MediaIoBaseDownload(f, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()
        return True
    except Exception as e:
        st.error(f"Erro ao baixar arquivo: {str(e)}")
        return False

def list_drive_files(service, folder_id, query_extra=""):
    """Lista arquivos em uma pasta específica do Google Drive"""
    try:
        query = f"'{folder_id}' in parents and trashed=false"
        if query_extra:
            query += f" and {query_extra}"
            
        results = service.files().list(
            q=query,
            spaces='drive',
            fields='files(id, name, mimeType)'
        ).execute()
        return results.get('files', [])
    except Exception as e:
        st.error(f"Erro ao listar arquivos: {str(e)}")
        return []

def find_vector_stores_in_drive(service, folder_id):
    """Encontra bases vetoriais no Google Drive"""
    vector_bases = {}
    
    # Lista todas as pastas no diretório
    folders = list_drive_folders(service, folder_id)
    
    for folder in folders:
        # Verifica se a pasta contém arquivos index.faiss e index.pkl
        files = list_drive_files(service, folder['id'])
        file_names = [file['name'] for file in files]
        
        if 'index.faiss' in file_names and 'index.pkl' in file_names:
            vector_bases[folder['name']] = {
                'folder_id': folder['id'],
                'files': {file['name']: file['id'] for file in files}
            }
            
    return vector_bases

# Funções para processamento de documentos  
def get_available_bases(base_dir):
    """Lista todas as bases de conhecimento disponíveis no diretório escolhido"""
    bases = {}
    if os.path.exists(base_dir):
        for base_name in os.listdir(base_dir):
            base_path = os.path.join(base_dir, base_name)
            if os.path.isdir(base_path) and os.path.exists(os.path.join(base_path, "index.faiss")):
                bases[base_name] = base_path
    return bases

def process_text_file(file_content: str) -> List[Document]:
    """Process plain text content"""
    if not file_content.strip():
        raise ValueError("O arquivo de texto está vazio")
    return [Document(page_content=file_content)]

def process_docx(file_content: bytes) -> List[Document]:
    """Process DOCX file content"""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        if not text.strip():
            raise ValueError("O documento Word está vazio")
        return [Document(page_content=text)]
    except Exception as e:
        raise ValueError(f"Erro ao processar arquivo DOCX: {str(e)}")

def process_pdf(file_content: bytes) -> List[Document]:
    """Process PDF file content"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = "\n".join([page.extract_text() for page in pdf_reader.pages])
        if not text.strip():
            raise ValueError("O arquivo PDF está vazio ou não contém texto extraível")
        return [Document(page_content=text)]
    except Exception as e:
        raise ValueError(f"Erro ao processar arquivo PDF: {str(e)}")

def process_html(file_content: str) -> List[Document]:
    """Process HTML file content"""
    try:
        soup = BeautifulSoup(file_content, 'html.parser')
        text = soup.get_text(separator="\n", strip=True)
        if not text:
            raise ValueError("O arquivo HTML está vazio ou não contém texto")
        return [Document(page_content=text)]
    except Exception as e:
        raise ValueError(f"Erro ao processar arquivo HTML: {str(e)}")

def process_csv(file_content: bytes) -> List[Document]:
    """Process CSV file content"""
    try:
        df = pd.read_csv(io.BytesIO(file_content))
        documents = []
        
        required_columns = ['id', 'assunto', 'classe', 'texto']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            raise ValueError(f"Colunas obrigatórias ausentes no CSV: {', '.join(missing_columns)}")
        
        df = df.fillna('')
        
        for _, row in df.iterrows():
            content = (
                f"ID: {row['id']}\n"
                f"Assunto: {row['assunto']}\n"
                f"Classe: {row['classe']}\n"
                f"Texto: {row['texto']}"
            )
            
            doc = Document(
                page_content=content,
                metadata={
                    'id': str(row['id']),
                    'assunto': str(row['assunto']),
                    'classe': str(row['classe'])
                }
            )
            documents.append(doc)
            
        if not documents:
            raise ValueError("Nenhum documento válido foi encontrado no CSV")
            
        return documents
        
    except pd.errors.EmptyDataError:
        raise ValueError("O arquivo CSV está vazio")
    except Exception as e:
        raise ValueError(f"Erro ao processar o CSV: {str(e)}")

@st.cache_data(show_spinner=False)
def process_uploaded_file(uploaded_file) -> List[Document]:
    """Process uploaded file based on its type"""
    try:
        file_type = uploaded_file.type
        file_content = uploaded_file.getvalue()
        
        if not file_content:
            raise ValueError("O arquivo está vazio")
        
        if file_type == "text/csv":
            return process_csv(file_content)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return process_docx(file_content)
        elif file_type == "application/pdf":
            return process_pdf(file_content)
        elif file_type == "text/html":
            return process_html(file_content.decode('utf-8'))
        elif file_type == "text/plain":
            return process_text_file(file_content.decode('utf-8'))
        else:
            raise ValueError(f"Tipo de arquivo não suportado: {file_type}")
            
    except Exception as e:
        st.error(f"Erro ao processar o arquivo: {str(e)}")
        raise

@st.cache_resource(show_spinner=False)
def create_vector_store(documents, vector_db_path):
    """Create and save vector store"""
    try:
        os.makedirs(vector_db_path, exist_ok=True)
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=200,
        )
        
        chunks = text_splitter.split_documents(documents)
        
        with st.spinner("Criando base de conhecimento..."):
            embedding_model = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
            vector_store = FAISS.from_documents(chunks, embedding_model)
            vector_store.save_local(vector_db_path)
            
        # Verificação final
        if not os.path.exists(os.path.join(vector_db_path, "index.faiss")):
            raise FileNotFoundError("Falha ao salvar vector store")
            
        return vector_store
        
    except Exception as e:
        if os.path.exists(vector_db_path):
            shutil.rmtree(vector_db_path)
        raise e

@st.cache_resource(show_spinner=False)
def load_vector_store(vector_db_path):
    """Load vector store from disk"""
    try:
        embedding_model = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        return FAISS.load_local(
            vector_db_path, 
            embedding_model,
            allow_dangerous_deserialization=True
        )
    except Exception as e:
        st.error(f"Erro ao carregar vector store: {str(e)}")
        return None

def main():
    """Função principal do aplicativo"""
    # Inicialização da API Google
    try:
        load_dotenv()
        google_api_key = os.getenv("google_api_key")
        
        if not google_api_key:
            st.error("Chave de API do Google não encontrada. Verifique seu arquivo .env")
            st.info("Você pode adicionar sua chave de API diretamente:")
            provided_key = st.text_input("Google API Key:", type="password")
            if provided_key:
                google_api_key = provided_key
            else:
                return
            
        genai.configure(api_key=google_api_key)
        model = genai.GenerativeModel(
            model_name="gemini-2.0-flash-thinking-exp-01-21", 
            generation_config={
                "temperature": 0.7,
                "top_p": 0.95,
                "top_k": 64,
                "max_output_tokens": 8192,
            }
        )
    except Exception as e:
        st.error(f"Erro ao configurar API Google: {str(e)}")
        logger.exception("Falha na inicialização da API:")
        return
    
    # Layout principal
    st.title("RAG App com Google Drive")
    
    # Conecta com Google Drive se necessário
    if not st.session_state.get('drive_service'):
        with st.expander("Configurar acesso ao Google Drive", expanded=True):
            st.session_state['drive_service'] = authorize_drive()
            if st.session_state['drive_service']:
                st.success("Conectado ao Google Drive com sucesso!")
    
    # Verifica se o serviço está disponível
    drive_available = st.session_state.get('drive_service') is not None
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        # Seletor de modo
        mode = st.radio(
            "Modo:",
            ["Criar base", "Usar base do Drive", "Carregar base local", "Deletar base"]
        )
        
        if mode == "Criar base":
            # Interface para criar base
            st.subheader("Criar nova base de conhecimento")
            
            if drive_available:
                st.info("A base será salva localmente. Para usar no Drive, crie e depois faça upload.")
                
            folder_path = st.text_input(
                "Pasta para salvar:",
                value=st.session_state['temp_dir'],
                help="Caminho para a pasta onde a base será salva"
            )
            
            if folder_path and not os.path.exists(folder_path):
                st.warning(f"A pasta {folder_path} será criada ao salvar")
            
            uploaded_file = st.file_uploader(
                "Upload do arquivo",
                type=['csv', 'docx', 'pdf', 'txt', 'html']
            )
            
            if uploaded_file and folder_path:
                base_name = st.text_input(
                    "Nome da base:",
                    help="Nome descritivo para a base"
                )
                
                if st.button("Criar Base", disabled=not base_name):
                    try:
                        with st.spinner("Processando arquivo..."):
                            base_name = "".join(c for c in base_name.strip() if c.isalnum() or c in ('-', '_'))
                            os.makedirs(folder_path, exist_ok=True)
                            base_path = os.path.join(folder_path, base_name)
                            
                            if os.path.exists(base_path):
                                st.error("Uma base com este nome já existe")
                            else:
                                documents = process_uploaded_file(uploaded_file)
                                if documents:
                                    vector_store = create_vector_store(documents, base_path)
                                    st.session_state['vector_store'] = vector_store
                                    st.success(f"Base '{base_name}' criada com sucesso!")
                                    
                                    # Opção para fazer download dos arquivos da base
                                    zip_path = shutil.make_archive(
                                        os.path.join(st.session_state['temp_dir'], base_name),
                                        'zip',
                                        base_path
                                    )
                                    with open(zip_path, 'rb') as f:
                                        st.download_button(
                                            "Baixar base para uso futuro",
                                            f,
                                            f"{base_name}.zip",
                                            "application/zip"
                                        )
                    except Exception as e:
                        st.error(f"Erro ao criar base: {str(e)}")
                        logger.exception("Erro na criação da base:")
                        
        elif mode == "Usar base do Drive" and drive_available:
            # Interface para usar base do Google Drive
            st.subheader("Carregar base do Google Drive")
            
            with st.spinner("Buscando pastas no Google Drive..."):
                folders = list_drive_folders_recursive(st.session_state['drive_service'])
                
            if not folders:
                st.warning("Nenhuma pasta encontrada no Google Drive ou problema de permissões")
            else:
                selected_folder_name = st.selectbox(
                    "Selecione a pasta no Drive:",
                    options=list(folders.keys())
                )
                
                if selected_folder_name:
                    folder_id = folders[selected_folder_name]
                    with st.spinner("Buscando bases de conhecimento..."):
                        vector_bases = find_vector_stores_in_drive(
                            st.session_state['drive_service'],
                            folder_id
                        )
                    
                    if not vector_bases:
                        st.warning(f"Nenhuma base de conhecimento encontrada em '{selected_folder_name}'")
                        st.info("Uma base válida deve conter arquivos 'index.faiss' e 'index.pkl'")
                    else:
                        selected_base = st.selectbox(
                            "Selecione a base para carregar:",
                            options=list(vector_bases.keys())
                        )
                        
                        if st.button("Carregar Base do Drive"):
                            try:
                                with st.spinner("Baixando arquivos do Google Drive..."):
                                    # Cria diretório temporário para a base
                                    base_temp_dir = os.path.join(
                                        st.session_state['temp_dir'],
                                        selected_base
                                    )
                                    os.makedirs(base_temp_dir, exist_ok=True)
                                    
                                    # Baixa os arquivos necessários
                                    base_info = vector_bases[selected_base]
                                    for filename, file_id in base_info['files'].items():
                                        if filename in ['index.faiss', 'index.pkl']:
                                            destination = os.path.join(base_temp_dir, filename)
                                            success = download_file_from_drive(
                                                st.session_state['drive_service'],
                                                file_id,
                                                destination
                                            )
                                            if not success:
                                                raise Exception(f"Falha ao baixar {filename}")
                                    
                                    # Carrega o vector store
                                    vector_store = load_vector_store(base_temp_dir)
                                    if vector_store:
                                        st.session_state['vector_store'] = vector_store
                                        st.success(f"Base '{selected_base}' carregada do Google Drive!")
                                    else:
                                        raise Exception("Falha ao carregar a base")
                            except Exception as e:
                                st.error(f"Erro ao carregar base do Drive: {str(e)}")
                                logger.exception("Erro no carregamento do Drive:")
        
        elif mode == "Carregar base local":
            # Interface para carregar base local via upload
            st.subheader("Carregar base via upload")
            
            uploaded_faiss = st.file_uploader("Arquivo index.faiss", type=['faiss'])
            uploaded_pkl = st.file_uploader("Arquivo index.pkl", type=['pkl'])
            
            if uploaded_faiss and uploaded_pkl:
                if st.button("Carregar Base"):
                    try:
                        with st.spinner("Processando arquivos..."):
                            # Cria diretório temporário
                            base_temp_dir = os.path.join(
                                st.session_state['temp_dir'],
                                f"uploaded_base_{pd.Timestamp.now().strftime('%Y%m%d%H%M%S')}"
                            )
                            os.makedirs(base_temp_dir, exist_ok=True)
                            
                            # Salva os arquivos
                            with open(os.path.join(base_temp_dir, "index.faiss"), 'wb') as f:
                                f.write(uploaded_faiss.getvalue())
                            with open(os.path.join(base_temp_dir, "index.pkl"), 'wb') as f:
                                f.write(uploaded_pkl.getvalue())
                            
                            # Carrega o vector store
                            vector_store = load_vector_store(base_temp_dir)
                            if vector_store:
                                st.session_state['vector_store'] = vector_store
                                st.success("Base carregada com sucesso!")
                            else:
                                raise Exception("Falha ao carregar a base")
                    except Exception as e:
                        st.error(f"Erro ao carregar base: {str(e)}")
                        logger.exception("Erro no carregamento via upload:")
            
            # Opção alternativa: carregar arquivo ZIP
            st.divider()
            st.subheader("Ou carregue um arquivo ZIP da base")
            
            uploaded_zip = st.file_uploader("Arquivo ZIP da base", type=['zip'])
            if uploaded_zip:
                if st.button("Extrair e Carregar ZIP"):
                    try:
                        with st.spinner("Extraindo e carregando base..."):
                            # Cria diretório temporário
                            zip_temp_dir = os.path.join(
                                st.session_state['temp_dir'],
                                f"zip_base_{pd.Timestamp.now().strftime('%Y%m%d%H%M%S')}"
                            )
                            os.makedirs(zip_temp_dir, exist_ok=True)
                            
                            # Salva o ZIP
                            zip_path = os.path.join(zip_temp_dir, "base.zip")
                            with open(zip_path, 'wb') as f:
                                f.write(uploaded_zip.getvalue())
                            
                            # Extrai o ZIP
                            extracted_dir = os.path.join(zip_temp_dir, "extracted")
                            shutil.unpack_archive(zip_path, extracted_dir, 'zip')
                            
                            # Procura os arquivos necessários
                            for root, dirs, files in os.walk(extracted_dir):
                                if 'index.faiss' in files and 'index.pkl' in files:
                                    # Encontrou diretório com os arquivos necessários
                                    vector_store = load_vector_store(root)
                                    if vector_store:
                                        st.session_state['vector_store'] = vector_store
                                        st.success("Base ZIP carregada com sucesso!")
                                        break
                            else:
                                st.error("Não foi possível encontrar arquivos index.faiss e index.pkl no ZIP")
                    except Exception as e:
                        st.error(f"Erro ao processar ZIP: {str(e)}")
                        logger.exception("Erro no processamento do ZIP:")
                        
        elif mode == "Deletar base" and drive_available:
            st.subheader("Deletar base")
            st.warning("Funcionalidade não implementada para bases no Google Drive")
            st.info("O Google Drive não permite excluir arquivos através desta interface por segurança")
            
            # Opção para limpar sessão
            if st.button("Limpar base carregada na sessão"):
                if 'vector_store' in st.session_state:
                    del st.session_state['vector_store']
                    st.success("Base removida da sessão")
                    st.experimental_rerun()
        
        elif not drive_available and mode in ["Usar base do Drive", "Deletar base"]:
            st.error("Conexão com Google Drive necessária")
            st.info("Expanda a seção 'Configurar acesso ao Google Drive' acima")
    
    # Coluna para consultas
    with col2:
        st.subheader("Consultar Base de Conhecimento")
        
        if st.session_state.get('vector_store'):
            query = st.text_input("Sua pergunta:")
            
            if query:
                try:
                    with st.spinner("Processando consulta..."):
                        results = st.session_state['vector_store'].similarity_search(query, k=5)
                        context = "\n".join([doc.page_content for doc in results])
                        
                        prompt = f"""
                        Contexto: {context}
                        
                        Pergunta: {query}
                        
                        Por favor, responda usando apenas as informações do contexto.
                        Se a informação não estiver disponível, diga que não pode responder.
                        
                        Resposta:
                        """
                        
                        response = model.generate_content(prompt)
                        
                        st.markdown("### Resposta:")
                        st.write(response.text)
                        
                        with st.expander("Ver contexto"):
                            st.write(context)
                except Exception as e:
                    st.error(f"Erro na consulta: {str(e)}")
                    logger.exception("Erro durante consulta:")
        else:
            st.info("Selecione ou crie uma base primeiro")
            st.divider()
            st.write("Status: Aguardando seleção de base")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        st.error(f"Erro não tratado: {str(e)}")
        logger.exception("Erro fatal na aplicação:")
