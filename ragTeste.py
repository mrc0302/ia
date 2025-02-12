import streamlit as st
import google.generativeai as genai
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import os
from dotenv import load_dotenv
import pandas as pd
from langchain_core.documents import Document
from typing import List
import docx
import PyPDF2
from bs4 import BeautifulSoup
import io
import shutil
from uuid import uuid4
from time import sleep

def get_base_dir():
    if 'base_dir' not in st.session_state:
        st.session_state['base_dir'] = os.path.normpath(r"./pages/base de conhecimento/")
    return st.session_state['base_dir']

def set_base_dir(new_dir):
    st.session_state['base_dir'] = os.path.normpath(new_dir)
    os.makedirs(new_dir, exist_ok=True)

def initialize_session_state():
    if 'vector_store' not in st.session_state:
        st.session_state['vector_store'] = None
    if 'confirm_delete' not in st.session_state:
        st.session_state['confirm_delete'] = False

def get_available_bases():
    """Lista todas as bases de conhecimento disponíveis no diretório base"""
    bases = {}
    base_dir = get_base_dir()
    if os.path.exists(base_dir):
        for base_name in os.listdir(base_dir):
            base_path = os.path.join(base_dir, base_name)
            if os.path.isdir(base_path) and os.path.exists(os.path.join(base_path, "index.faiss")):
                bases[base_name] = base_path
    return bases

def process_text_file(file_content: str) -> List[Document]:
    """Process plain text content"""
    if not file_content.strip():
        raise ValueError("O arquivo de texto está vazio")
    return [Document(page_content=file_content)]

def process_docx(file_content: bytes) -> List[Document]:
    """Process DOCX file content"""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        if not text.strip():
            raise ValueError("O documento Word está vazio")
        return [Document(page_content=text)]
    except Exception as e:
        raise ValueError(f"Erro ao processar arquivo DOCX: {str(e)}")

def process_pdf(file_content: bytes) -> List[Document]:
    """Process PDF file content"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = "\n".join([page.extract_text() for page in pdf_reader.pages])
        if not text.strip():
            raise ValueError("O arquivo PDF está vazio ou não contém texto extraível")
        return [Document(page_content=text)]
    except Exception as e:
        raise ValueError(f"Erro ao processar arquivo PDF: {str(e)}")

def process_html(file_content: str) -> List[Document]:
    """Process HTML file content"""
    try:
        soup = BeautifulSoup(file_content, 'html.parser')
        text = soup.get_text(separator="\n", strip=True)
        if not text:
            raise ValueError("O arquivo HTML está vazio ou não contém texto")
        return [Document(page_content=text)]
    except Exception as e:
        raise ValueError(f"Erro ao processar arquivo HTML: {str(e)}")

def process_csv(file_content: bytes) -> List[Document]:
    """Process CSV file content"""
    try:
        df = pd.read_csv(io.BytesIO(file_content))
        documents = []
        
        required_columns = ['id', 'assunto', 'classe', 'texto']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            raise ValueError(f"Colunas obrigatórias ausentes no CSV: {', '.join(missing_columns)}")
        
        df = df.fillna('')
        
        for _, row in df.iterrows():
            content = (
                f"ID: {row['id']}\n"
                f"Assunto: {row['assunto']}\n"
                f"Classe: {row['classe']}\n"
                f"Texto: {row['texto']}"
            )
            
            doc = Document(
                page_content=content,
                metadata={
                    'id': str(row['id']),
                    'assunto': str(row['assunto']),
                    'classe': str(row['classe'])
                }
            )
            documents.append(doc)
            
        if not documents:
            raise ValueError("Nenhum documento válido foi encontrado no CSV")
            
        return documents
        
    except pd.errors.EmptyDataError:
        raise ValueError("O arquivo CSV está vazio")
    except Exception as e:
        raise ValueError(f"Erro ao processar o CSV: {str(e)}")

def process_uploaded_file(uploaded_file) -> List[Document]:
    """Process uploaded file based on its type"""
    try:
        st.write(f"Processando arquivo: {uploaded_file.name}")
        file_type = uploaded_file.type
        file_content = uploaded_file.getvalue()
        
        if not file_content:
            raise ValueError("O arquivo está vazio")
        
        st.write(f"Tamanho do arquivo: {len(file_content)} bytes")
        
        if file_type == "text/csv":
            return process_csv(file_content)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return process_docx(file_content)
        elif file_type == "application/pdf":
            return process_pdf(file_content)
        elif file_type == "text/html":
            return process_html(file_content.decode('utf-8'))
        elif file_type == "text/plain":
            return process_text_file(file_content.decode('utf-8'))
        else:
            raise ValueError(f"Tipo de arquivo não suportado: {file_type}")
            
    except Exception as e:
        st.error(f"Erro ao processar o arquivo: {str(e)}")
        raise

def create_vector_store(documents, vector_db_path):
    """Create and save vector store"""
    try:
        os.makedirs(vector_db_path, exist_ok=True)
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=200,
        )
        
        st.write("Dividindo documentos...")
        chunks = text_splitter.split_documents(documents)
        st.write(f"Total de chunks: {len(chunks)}")
        
        st.write("Criando vector store...")
        vector_store = FAISS.from_documents(chunks, embedding_model)
        
        st.write("Salvando vector store...")
        vector_store.save_local(vector_db_path)
        
        # Verificação final
        if not os.path.exists(os.path.join(vector_db_path, "index.faiss")):
            raise FileNotFoundError("Falha ao salvar vector store")
            
        return vector_store
        
    except Exception as e:
        st.error(f"Erro ao criar vector store: {str(e)}")
        if os.path.exists(vector_db_path):
            shutil.rmtree(vector_db_path)
        return None

def load_vector_store(vector_db_path):
    """Load vector store from disk"""
    try:
        return FAISS.load_local(
            vector_db_path, 
            embedding_model,
            allow_dangerous_deserialization=True
        )
    except Exception as e:
        st.error(f"Erro ao carregar vector store: {str(e)}")
        return None

def create_new_knowledge_base():
    """Interface para criar nova base de conhecimento"""
    new_base_dir = st.sidebar.text_input(
        "Diretório para base de conhecimento:",
        value=get_base_dir(),
        help="Caminho completo para o diretório onde as bases serão armazenadas"
    )
    
    if new_base_dir != get_base_dir():
        set_base_dir(new_base_dir)
    
    uploaded_file = st.sidebar.file_uploader(
        "Upload do arquivo",
        type=['csv', 'docx', 'pdf', 'txt', 'html']
    )
    
    if uploaded_file:
        base_name = st.sidebar.text_input(
            "Nome da base:",
            help="Nome descritivo para a base"
        )
        
        if base_name and st.sidebar.button("Criar Base"):
            try:
                base_name = "".join(c for c in base_name.strip() if c.isalnum() or c in ('-', '_'))
                base_path = os.path.join(get_base_dir(), base_name)
                
                if os.path.exists(base_path):
                    st.sidebar.error("Uma base com este nome já existe")
                    return
                
                documents = process_uploaded_file(uploaded_file)
                if not documents:
                    raise ValueError("Nenhum documento processado")
                    
                vector_store = create_vector_store(documents, base_path)
                if vector_store:
                    st.session_state['vector_store'] = vector_store
                    st.success(f"Base '{base_name}' criada com sucesso!")
                    st.rerun()
                    
            except Exception as e:
                st.error(f"Erro ao criar base: {str(e)}")

def use_existing_bases():
    """Interface para usar bases existentes"""
    new_base_dir = st.sidebar.text_input(
        "Diretório das bases de conhecimento:",
        value=get_base_dir(),
        help="Caminho completo para o diretório onde as bases estão armazenadas"
    )
    
    if new_base_dir != get_base_dir():
        set_base_dir(new_base_dir)
    
    bases = get_available_bases()
    
    if not bases:
        st.sidebar.warning(f"Nenhuma base encontrada em {get_base_dir()}")
        return
        
    selected_bases = st.sidebar.multiselect(
        "Selecione as bases:",
        options=list(bases.keys())
    )
    
    if selected_bases and st.sidebar.button("Carregar Bases"):
        try:
            combined_store = None
            
            for base_name in selected_bases:
                vector_store = load_vector_store(bases[base_name])
                
                if vector_store:
                    if combined_store is None:
                        combined_store = vector_store
                    else:
                        break
            
            if combined_store:
                st.session_state['vector_store'] = combined_store
                st.success("Base(s) carregada(s) com sucesso!")
                st.rerun()
                
        except Exception as e:
            st.error(f"Erro ao carregar bases: {str(e)}")

def delete_knowledge_base():
    """Interface para deletar bases"""
    new_base_dir = st.sidebar.text_input(
        "Diretório das bases de conhecimento:",
        value=get_base_dir(),
        help="Caminho completo para o diretório onde as bases estão armazenadas"
    )
    
    if new_base_dir != get_base_dir():
        set_base_dir(new_base_dir)
    
    bases = get_available_bases()
    
    if not bases:
        st.sidebar.warning(f"Nenhuma base para deletar em {get_base_dir()}")
        return
        
    base_to_delete = st.sidebar.selectbox(
        "Selecione a base para excluir:",
        options=list(bases.keys())
    )
    
    if st.sidebar.button("Deletar Base"):
        try:
            base_path = bases[base_to_delete]
            shutil.rmtree(base_path)
            st.success(f"Base '{base_to_delete}' deletada")
            st.rerun()
        except Exception as e:
            st.error(f"Erro ao deletar base: {str(e)}")

def handle_query():
    """Interface para consultas"""
    if st.session_state.get('vector_store'):
        query = st.text_input("Sua pergunta:")
        
        if query:
            try:
                results = st.session_state['vector_store'].similarity_search(query, k=5)
                
                context = "\n".join([doc.page_content for doc in results])
                
                prompt = f"""
                Contexto: {context}
                
                Pergunta: {query}
                
                Por favor, responda usando apenas as informações do contexto.
                Se a informação não estiver disponível, diga que não pode responder.
                
                Resposta:
                """
                
                response = model.generate_content(prompt)
                
                st.subheader("Resposta:")
                st.write(response.text)
                
                with st.expander("Ver contexto"):
                    st.write(context)
                    
            except Exception as e:
                st.error(f"Erro na consulta: {str(e)}")
    else:
        st.info("Selecione ou crie uma base primeiro")

# Configuração principal da aplicação
#st.set_page_config(page_title="RAG App", layout="wide")
st.title("RAG App com Google Gemini")

initialize_session_state()

# Configuração da API
load_dotenv()
google_api_key = os.getenv("google_api_key")

genai.configure(api_key=google_api_key)

generation_config = {
    "temperature": 0.7,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 65536,
}

# Modelos globais
global embedding_model, model
embedding_model = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
model = genai.GenerativeModel(
    model_name="gemini-2.0-flash-thinking-exp-01-21",
    generation_config=generation_config,
)

# Interface principal
st.sidebar.header("Opções")
mode = st.sidebar.radio(
    "Modo:",
    ["Criar base", "Usar base existente", "Deletar base"]
)

if mode == "Criar base":
    create_new_knowledge_base()
elif mode == "Usar base existente":
    use_existing_bases()
else:
    delete_knowledge_base()

handle_query()
