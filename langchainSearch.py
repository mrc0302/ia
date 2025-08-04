
#import asyncio
#import nest_asyncio
import time
from langchain.text_splitter import RecursiveCharacterTextSplitter
import re
from PyPDF2 import PdfReader
import sqlite3
import csv
import docx
from googlesearch import search
from serpapi import GoogleSearch
import streamlit as st
from langchain_community.vectorstores import FAISS
import google.generativeai as genai
from langchain_google_genai import GoogleGenerativeAI
from langchain.prompts import PromptTemplate
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.retrievers import MultiQueryRetriever
from langchain.chains.query_constructor.schema import AttributeInfo
from langchain.retrievers.self_query.base import SelfQueryRetriever
from langchain.chains.question_answering import load_qa_chain
import os
from dotenv import load_dotenv
import json
import pandas as pd
from io import StringIO
from langchain.docstore.document import Document
from langchain.prompts.chat import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.schema import Document
from langchain.chains import create_retrieval_chain

def main():

    # Configuração da página
    #st.set_page_config(
     #   page_title="Sistema de Modelos Judiciais",
      #  page_icon="🧊", 
       # layout="wide",  
        #initial_sidebar_state="expanded"
    #)

    #def load_css(file_name):
     #   with open(file_name, encoding='utf-8') as f:  # Adicione encoding='utf-8'
      #      st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

    #load_css("static/styles.css")  

        # Carregar variáveis de ambiente
    load_dotenv()
    google_api_key = os.getenv("google_api_key")
    #-------------------------------------------------------------------------------
    genai.configure(api_key=os.getenv("google_api_key"))  

    # Modelos disponíveis
    models = [
        "gemini-1.5-flash",
        "gemini-2.0-flash",              # Mais rápido
        "learnlm-2.0-flash-experimental", # Seu modelo atual
        "gemini-2.0-flash-thinking-exp-01-21",
        "gemma-3n-e2b-it",
        "gemma-3n-e4b-it",
        "gemma-3-1b-it",
        "gemma-3-4b-it",
        "gemma-3-12b-it",
        "gemma-3-27b-it",
        "gemini-2.0-flash-preview-image-generation",
        "gemini-2.0-flash-lite",
        "gemini-2.5-pro",
        "gemini-2.5-flash",
        "gemini-2.5-flash-lite",
        

        
    ]
    # Tipos MIME suportados
    mime_types = [
        "text/plain",     # Texto simples (padrão)
        "text/html",      # HTML
        "application/json", # JSON
        "text/markdown"   # Markdown
    ]
    # Create the model
    generation_config = {
        "temperature": 0.7,
        "top_p": 0.95,
        "top_k": 64,
        "max_output_tokens": 8024,
        "response_mime_type":  "text/plain" ,
        
        
   }
    model_name="learnlm-2.0-flash-experimental"
    model = genai.GenerativeModel(model_name=model_name)
    
    llm2 = GoogleGenerativeAI(
            model=model_name,
            google_api_key=google_api_key,
            temperature=0.7
        )
    
    llm3 = ChatGoogleGenerativeAI(
        model=model_name,
        temperature=0,
        google_api_key=google_api_key        
    )

    def get_mq_retriever(vector_store, llm):
                
        # Retriever básico
        basic_retriever = vector_store.as_retriever(search_kwargs={"k": 3})
        
        # MultiQueryRetriever
        multi_query_retriever = MultiQueryRetriever.from_llm(
            retriever=basic_retriever,
            llm=llm
        )

        return multi_query_retriever

    def get_embeddings():

        embeddings = GoogleGenerativeAIEmbeddings(
                model="models/embedding-001",
                google_api_key=google_api_key
            )
        
        return embeddings

    # Initialize session state for chat history
    if "messages" not in st.session_state:
        st.session_state.messages = []

    def limpar_tudo():
       # init_session()

        st.session_state.messages = []
        st.session_state.chat_history = []
        st.session_state.documentos_contexto = []
        st.session_state.uploaded_files = []
        st.session_state.use_web_search = False
        st.session_state.documentos_contexto = []               
        st.session_state.chat_session = model.start_chat(history=[])
        st.session_state.campo_selecionado = " "
        st.session_state.valor_campo_selecionado = " "
        st.session_state.texto_livre_input = ""
        

         # Limpar qualquer chave relacionada aos controles
        keys_to_clear = ['texto_livre_key', 'campo_key', 'valor_campo_key']
        for key in keys_to_clear:
            if key in st.session_state:
                del st.session_state[key]
        
        # Força o recarregamento da página
        st.rerun()

    def limpar_apenas_arquivos():
        """Limpa apenas os arquivos carregados, mantendo o chat"""
        st.session_state.uploaded_files = []

    def converter_sqlite_para_vectorstore():
        """
        Função integrada para converter SQLite para FAISS Vector Store
        """
        try:
            import sqlite3
            
            # Parâmetros
            db_path = "bdProjetoForense.db"
            vectordb_path = "vectordb"
            
            if not os.path.exists(db_path):
                st.error(f"Arquivo do banco de dados não encontrado: {db_path}")
                return None
            
            with st.spinner("Conectando ao banco de dados..."):
                # Conectar ao SQLite
                conn = sqlite3.connect(db_path)
                cursor = conn.cursor()
                
                # Buscar dados
                cursor.execute("SELECT id, assunto, classe, texto, img, html FROM tbAssunto")
                rows = cursor.fetchall()
                conn.close()
            
            if not rows:
                st.error("Nenhum dado encontrado na tabela 'tbAssunto'")
                return None
            
            st.success(f"Encontrados {len(rows)} registros")
            
            with st.spinner("Criando documentos..."):
                documents = []
                for row in rows:
                    id_doc, assunto, classe, texto, img, html = row
                    
                    # Combinar conteúdo
                    content_parts = []
                    if texto:
                        content_parts.append(texto)
                    if html:
                        content_parts.append(html)
                    
                    if not content_parts:
                        continue
                    
                    page_content = "\n\n".join(content_parts)
                    
                    metadata = {
                        "id": id_doc,
                        "assunto": assunto or "Não especificado",
                        "classe": classe or "Não especificado", 
                        "img": img or "",
                        "texto": texto or ""
                    }
                    
                    doc = Document(page_content=page_content, metadata=metadata)
                    documents.append(doc)
            
            if not documents:
                st.error("Nenhum documento válido criado")
                return None
            
            with st.spinner("Dividindo documentos em chunks..."):
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=800,
                    chunk_overlap=200,
                    length_function=len,
                    separators=["\n\n", "\n", ".", " ", ""]
                )
                chunks = text_splitter.split_documents(documents)
            
            with st.spinner("Criando vector store FAISS..."):
                embeddings = get_embeddings()
                vector_store = FAISS.from_documents(chunks, embeddings)
                
                # Salvar
                os.makedirs(vectordb_path, exist_ok=True)
                vector_store.save_local(vectordb_path)
            
            st.success(f"✅ Vector store criado com sucesso!")
            st.info(f"📁 Salvo em: {os.path.abspath(vectordb_path)}")
            st.info(f"📊 Documentos: {len(documents)} | Chunks: {len(chunks)}")
            
            return vector_store
        
        except Exception as e:
            st.error(f"Erro na conversão: {str(e)}")
            return None
        
    def extract_text_from_pdf(pdf_file):
        pdf_reader = PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text

    def extract_text_from_docx(docx_file):
        doc = docx.Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def process_uploaded_file(uploaded_file):
        if uploaded_file.type == "application/pdf":
            return {
                'name': uploaded_file.name,
                'content': extract_text_from_pdf(uploaded_file),
                'type': 'pdf'
            }
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return {
                'name': uploaded_file.name,
                'content': extract_text_from_docx(uploaded_file),
                'type': 'docx'
            }
        # elif uploaded_file.type == "text/csv":
        #     return {
        #         'name': uploaded_file.name,
        #         'content': extract_text_from_csv(uploaded_file),
        #         'type': 'csv'
        #     }
        else:
            raise ValueError("Formato de arquivo não suportado")

    def split_documents(documents):
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,  # tamanho do chunk
            chunk_overlap=200,  # sobreposição
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        chunks = text_splitter.split_documents(documents)
        
        return chunks
   
    def create_faiss_vectorstore(chunks):
        
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=google_api_key
        )        
        vectorstore = FAISS.from_documents(
            documents=chunks,
            embedding=embeddings
        )
        
        return vectorstore
   
    @st.cache_resource
    def carregar_vector_store():
        try:
            # Tentar obter o loop atual ou criar um novo
            try:
                loop = asyncio.get_event_loop()
            except RuntimeError:
                # Se não há loop, criar um novo
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
            
            # Permitir loops aninhados
            nest_asyncio.apply()
            
            embeddings = GoogleGenerativeAIEmbeddings(
                model="models/embedding-001",
                google_api_key=google_api_key
            )
            return FAISS.load_local(
                "vectordb", 
                embeddings, 
                allow_dangerous_deserialization=True
            )
        except Exception as e:
            st.error(f"Erro ao carregar vector store: {str(e)}")
            return None
        
    vector_store = carregar_vector_store()

    def init_session():
        # Inicializar estado da sessão
        if "messages" not in st.session_state:
            st.session_state.messages = []

        if 'documentos_contexto' not in st.session_state:
            st.session_state.documentos_contexto = []

        if 'chat_history' not in st.session_state:
            st.session_state.chat_history = []

        # Inicializar estado da sessão
        if 'uploaded_files' not in st.session_state:
            st.session_state.uploaded_files = []

        if 'use_web_search' not in st.session_state:
            st.session_state.use_web_search = False

        if 'use_juris_search' not in st.session_state:
            st.session_state.use_juris_search = False

        if "chat_session" not in st.session_state: #importante
            st.session_state.chat_session = model.start_chat(history=[])    

        if 'prompt' not in st.session_state:        
            st.session_state.prompt = None    
 
        if 'campo_selecionado' not in st.session_state:
           st.session_state.campo_selecionado = " "
    
        if 'valor_campo_selecionado' not in st.session_state:
            st.session_state.valor_campo_selecionado = " "

        if 'texto_livre_input' not in st.session_state:
            st.session_state.texto_livre_input = ""

        if 'model_selecionado' not in st.session_state: 
            st.session_state.model_selecionado = model_name
        
        if 'on_sqlite_troggle' not in st.session_state:
            st.session_state.on_sqlite_troggle = False
        
        # Configurações de busca
        if 'search_config' not in st.session_state:
            st.session_state.search_config = {
                'search_type': 'mmr',
                'k': 10,
                'fetch_k': 20,
                'lambda_mult': 0.5,
                'score_threshold': 0.4
            }
        
        # Configurações de geração
        if 'generation_config' not in st.session_state:
            st.session_state.generation_config = {
                'temperature': 0.1,
                'top_p': 0.9,
                'top_k': 40,
                'max_output_tokens': 8192,
                'candidate_count': 1,
                'stop_sequences': []  # Lista vazia por padrão
            }

    init_session()

    def perform_web_search(query, num_results=5):
        try:
            # Configurar a busca
            params = {
                "engine": "google",
                "q": query,
                "api_key": os.getenv("SERPAPI_KEY"),
                "num": num_results,
                "hl": "pt-br",
                "gl": "br"
            }
            
            search = GoogleSearch(params)
            results = search.get_dict()
            
            # Extrair e formatar resultados
            if "organic_results" in results:
                formatted_results = []
                for result in results["organic_results"][:num_results]:
                    title = result.get("title", "")
                    snippet = result.get("snippet", "")
                    link = result.get("link", "")
                    formatted_results.append(
                        f"Título: {title}\n"
                        f"Resumo: {snippet}\n"
                        f"Link: {link}"
                    )
                
                return "\n\n".join(formatted_results)
            else:
                return "Nenhum resultado encontrado."
                
        except Exception as e:
            print(f"Erro na busca: {str(e)}")  # Para debug
            return "Não foi possível realizar a busca no momento."

    def extrair_campos_unicos(vector_store):
        try:
            resultados = vector_store.similarity_search("", k=100000)
            classes = set()
            assuntos = set()
            for doc in resultados:
                metadata = getattr(doc, 'metadata', {})
                if metadata.get('classe'):
                    classes.add(metadata['classe'])
                if metadata.get('assunto'):
                    assuntos.add(metadata['assunto'])
            return sorted(list(classes)), sorted(list(assuntos))
        except Exception as e:
            st.error(f"Erro ao extrair campos: {str(e)}")
            return [], []
    
    def create_vector_store(documents, vector_db_path):
        """Create and save vector store"""
        try:
            os.makedirs(vector_db_path, exist_ok=True)
            
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=800,
                chunk_overlap=200,
            )
            
            st.write("Dividindo documentos...")
            chunks = text_splitter.split_documents(documents)
            st.write(f"Total de chunks: {len(chunks)}")
            
            st.write("Criando vector store...")
            vector_store = FAISS.from_documents(chunks, get_embeddings)
                      
                
            return vector_store
            
        except Exception as e:
            st.error(f"Erro ao criar vector store: {str(e)}")
            
            return None

    llm = ChatGoogleGenerativeAI(
    model="gemma-3n-e2b-it",
    temperature=0,
    google_api_key=google_api_key,
    
)

    def limpar_query(query_suja):
        query_limpa = query_suja.strip()
        query_limpa = query_limpa.replace('```sqlite', '').replace('```sql', '').replace('```', '')
        query_limpa = query_limpa.strip()
        return query_limpa


    def searchSQL(pergunta):      

        # Criar gerador de query
        #write_query = create_sql_query_chain(llm, db)
        
        llm = ChatGoogleGenerativeAI(
            model="gemma-3n-e2b-it",
            temperature=0,
            google_api_key=google_api_key,
            
        )
        prompt_template = """
                Você é um {role}.
                Tarefa: {task}
                Contexto: {context}                
                Pergunta: {question}
            """

        prompt = prompt_template.format(
            role="Você é um assistente especialista em gerar instruções SQL em banco de dados Sqlite",
            task="""converta em a pergunta do usuário em uma instrução sql para o banco de dados, fornecendo apenas a query SQL, sem explicações adicionais. Tabela : tbAssunto Colunas: id, Classe, Assunto, Texto""",
            question=pergunta,
            context=""
        )

        conn = sqlite3.connect("bdProjetoForense.db")

        conn.row_factory = sqlite3.Row  # ✨ Esta linha é a solução!
        
        cursor = conn.cursor()

        # # Chamar o llm diretamente com uma string
        llm_response = llm.invoke(prompt)
        print()
        print(llm_response.content) # Use .content para obter o texto da resposta

        resultado= limpar_query(llm_response.content)


        cursor.execute(resultado)  # Executar a query gerada

        results = cursor.fetchall()  # Obter todos os resultados

        metadados = []  # Lista para armazenar metadados
        
        print(len(results))
        
        if results:            
            for i, col in enumerate(results):

               print(f"Índice {i}: {col[0]}")
               metadado_src = (f"""[id : {col[0]}  - CLASSE  {col[1]} - ASSUNTO : {col[2]}
                                TEXTO : {col[3]}""")
               metadados.append(metadado_src)
        
        return metadados  
          
    def gerar_resposta(query, vector_store, contexto_docs, historico, uploaded_files, campo, valor_campo, model_name):
    
        try:

            llm_ = GoogleGenerativeAI(
            model=st.session_state.model_selecionado,
            google_api_key=google_api_key,
            temperature=0.7,
            generation_config=generation_config,
            
            )   
            contexto_completo = " "            
            retrieved_docs = []
            arquivos_texto = []
            metadados = []  # Lista para armazenar metadados
            on_sqlite_troggle= st.session_state.on_sqlite_troggle

            if contexto_docs: # são os documentos da consultas                
                retrieved_docs = contexto_docs
                
                for doc in retrieved_docs:
                    metadata_str = f"[Classe: {doc.metadata.get('classe', 'N/A')}, Assunto: {doc.metadata.get('assunto', 'N/A')}]"
                    arquivos_texto.append(f"{metadata_str}\n{doc.page_content}") # além dos metadados inclui o conteúdo dos documentos
                    metadados.append(doc.metadata)  # Adiciona metadados à lista de metadados com as informações dos documentos
            
            elif on_sqlite_troggle:

                arquivos_texto= searchSQL(query)
                
            elif uploaded_files:                   

                for doc in uploaded_files:
                    doc_obj = Document(
                        page_content=doc['content'], 
                        metadata={'name': doc['name'], 'type': doc['type']}
                    )
                    retrieved_docs.append(doc_obj) # o retrieved_docs será utilizado na resposta como parâmetro

                    for doc in retrieved_docs:
                        metadata_str = f"[Classe: {doc.metadata.get('classe', 'N/A')}, Assunto: {doc.metadata.get('assunto', 'N/A')}]"
                        arquivos_texto.append(f"{metadata_str}\n{doc.page_content}") # será utilizado no contexto não se é o melhor
                        metadados.append(doc.metadata)  # Adiciona metadados à lista                                                                           
        
            elif campo == "classe" and valor_campo:
                # USAR CONFIGURAÇÕES PERSONALIZADAS DO SESSION STATE
                search_kwargs = {
                    "k": st.session_state.search_config['k'],
                    "fetch_k": st.session_state.search_config['fetch_k'],
                    "lambda_mult": st.session_state.search_config['lambda_mult'],
                    "score_threshold": st.session_state.search_config['score_threshold'],
                    "filter": {campo: valor_campo}
                }                
                retriever = vector_store.as_retriever(
                    search_type=st.session_state.search_config['search_type'], 
                    search_kwargs=search_kwargs
                )
                retrieved_docs = retriever.invoke(query)
                for doc in retrieved_docs:
                    metadata_str = f"[Classe: {doc.metadata.get('classe', 'N/A')}, Assunto: {doc.metadata.get('assunto', 'N/A')}]"
                    arquivos_texto.append(f"{metadata_str}\n{doc.page_content}")
                    metadados.append(doc.metadata)  # Adiciona metadados à lista    

            else:
               
               
                mq_retriever = MultiQueryRetriever.from_llm(
                            retriever=vector_store.as_retriever(search_kwargs={"k": 10}),
                            llm=llm_
                )         

                arquivos_texto = [] 
                metadados = []  # Lista para armazenar os metadados dos documentos
                # Recuperar documentos relevantes
                retrieved_docs = mq_retriever.get_relevant_documents(query=query, k=10)

                for doc in retrieved_docs:
                    

                    metadata_str = (f"""[id : {doc.metadata.get('id', 'N/A')}  - CLASSE  {doc.metadata.get('classe', 'N/A')} - ASSUNTO : {doc.metadata.get('assunto', 'N /A')}
                                TEXTO : {doc.metadata.get('texto', 'N /A')} - HTML {doc.metadata.get('html', 'N/A')}""")
                    
                    arquivos_texto.append(f"{metadata_str}") # além dos metadados inclui o conteúdo dos documentos
                # metadados.append(doc.metadata)  # Adiciona metadados à lista de metadados com as informações dos documentos
                
                #contexto_completo = "\n\n".join([doc.page_content for doc in retrieved_docs])
            contexto_completo = "\n\n".join(arquivos_texto)
                                    
            prompt_template = """
                Você é um {role}.
                Tarefa: {task}
                Contexto: {context}
                output_format: {output_format}
                só utilize o histórico:{history} se for pedido ou necessário para resposta
                Pergunta: {question}
            """

            prompt = prompt_template.format(
                role="Você é um assistente especialista em direito",
                task="""Responda as perguntas do usuário sempre em português de forma clara COM BASE NO CONTEXTO. SEMPRE forneça A DESCRIÇÃO DO ASSUNTO entre []""",
                context=contexto_completo,
                output_format="formate TEXTO PLANO . não use html", 
                history=historico,
                question=query
            ) 
                     
            response = llm_.invoke(prompt)
            
            return str(response)  # Trabalha com o modelo de geração de texto
            
        except Exception as e:
            return f"Erro ao gerar resposta: {str(e)}"
    
   
    def busca_combinada(vector_store, campo, valor_campo, texto_livre, num_results):
        
        try:

            resultados = []

            total_vectors = vector_store.index.ntotal #  # Total de registros no vector store
             # Primeiro busca pelo texto livre se não houver filtro de metadados selecionado
            todos_docs = vector_store.similarity_search("", k=total_vectors)
            
            if not (campo =="classe" or campo =="assunto") and texto_livre:               
                              
                termo_lower = texto_livre.lower() # termo de busca em minúsculas
                
                for doc in todos_docs:      # doc é um objeto Document com page_content e metadata
                    
                    valor_campo = doc.metadata.get("texto", "").lower() # Se "texto" existe: retorna o valor
                                                                        # Se "texto" NÃO existe: retorna "" (string vazia)
                    
                    # Aqui está a "busca parcial" - equivalente ao LIKE '%termo%'
                    if termo_lower in valor_campo:
                        resultados.append(doc) #adiciono a lista de resultados com metadados
                         
                        if len(resultados) >= num_results:
                            break                   
            
            elif campo =="classe" and  valor_campo and texto_livre:  # Se houver filtro de metadados
                
                termo_lower = texto_livre.lower() # termo de busca em minúsculas
                
                for doc in todos_docs:      # doc é um objeto Document com page_content e metadata
                    
                    classe_content = doc.metadata.get("classe", "") 

                    if classe_content == valor_campo and termo_lower in doc.page_content.lower():# se a classe do documento é igual ao valor do campo 
                                                                                            # e o termo de busca está no conteúdo   

                        resultados.append(doc) #adiciono a lista de resultados com metadados
                        
                        if len(resultados) >= num_results:
                            break    
         
            elif campo =="assunto" and  valor_campo and texto_livre:  # Se houver filtro de metadados
                
                termo_lower = texto_livre.lower() # termo de busca em minúsculas
                
                for doc in todos_docs:      # doc é um objeto Document com page_content e metadata
                    
                    assunto_content = doc.metadata.get("assunto", "") 

                    if assunto_content == valor_campo and termo_lower in doc.page_content.lower():# se a assunto do documento é igual ao valor do campo 
                                                                                            # e o termo de busca está no conteúdo   

                        resultados.append(doc) #adiciono a lista de resultados com metadados
                        
                        if len(resultados) >= num_results:
                            break         

            elif campo  and  valor_campo and not texto_livre:  # Se houver filtro de metadados
                
                for doc in todos_docs:
                    classe_content = doc.metadata.get("classe", "") 

                    if classe_content == valor_campo:
                        resultados.append(doc) #adiciono a lista de resultados com metadados            

            st.write(f"Total de documentos encontrados: {len(resultados)}")
            return resultados[:num_results]  # retorna apenas os primeiros num_results documentos
            
        except Exception as e:
            st.error(f"Erro na busca: {str(e)}")
            print(f"Erro detalhado: {str(e)}")  # Para debug
            return [] 
   
    def extrair_campos_unicos(vector_store):
        try:
            resultados = vector_store.similarity_search("", k=100000)
            classes = set()
            assuntos = set()
            for doc in resultados:
                metadata = getattr(doc, 'metadata', {})
                if metadata.get('classe'):
                    classes.add(metadata['classe'])
                if metadata.get('assunto'):
                    assuntos.add(metadata['assunto'])
            return sorted(list(classes)), sorted(list(assuntos))
        except Exception as e:
            st.error(f"Erro ao extrair campos: {str(e)}")
            return [], []

    def encontrar_urls(texto):
        padrao = r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+'
        urls = re.findall(padrao, texto)
        return urls

    def stream_data(resposta):
        for word in resposta.split(" "):
            yield word + " "
            time.sleep(0.02)

   
    if vector_store:
        # Sidebar
        colSideBAr, colChatbot = st.columns([1, 4])
        with colSideBAr:
        #with st.sidebar:            
            tab1, tab2, tab3 = st.tabs(["Pesquisa", "Arquivos","Configurações"])
            with tab1:                 
                expander_pesq = st.expander("🔍 Filtros de Busca", expanded=True)                
                with expander_pesq:            
                    texto_livre = st.text_input(
                        "Texto Livre:", 
                        placeholder="Ex: termo1, termo2",
                        value=st.session_state.texto_livre_input,  
                        key="texto_livre_key"  
                     )
                    # Atualizar o estado quando o valor mudar
                    st.session_state.texto_livre_input = texto_livre
                                 
                    campo = st.selectbox(
                        "Campo:", 
                        [" ", "classe", "assunto"], 
                        help="Filtrar por campo",
                        index=[" ", "classe", "assunto"].index(st.session_state.campo_selecionado),  
                        key="campo_key" 
                    )
                    st.session_state.campo_selecionado = campo

                    valores = [] 
                    classes, assuntos = extrair_campos_unicos(vector_store)               
                
                    if campo:                    
                        valores = classes if campo == "classe" else assuntos
                    
                    # Garantir que o valor selecionado seja válido para a lista atual
                    
                    if st.session_state.valor_campo_selecionado not in [" "] + valores:
                        st.session_state.valor_campo_selecionado = " "
                    
                    valor_campo = st.selectbox(
                        "Valor do Campo:",
                        [" "] + valores, 
                        help="Selecione o valor do campo",
                        index=([" "] + valores).index(st.session_state.valor_campo_selecionado),  
                        key="valor_campo_key"                                              
                    )

                    # Atualizar o estado quando o valor mudar
                    st.session_state.valor_campo_selecionado = valor_campo
                    
                    num_results = st.slider("Nº de Resultados:", 1, 20, 4)

                    col1, col2 = st.columns([1,1])
                    with col1:
                        buscar = st.button("Buscar")
                        if buscar:
                            if not any([texto_livre, campo and valor_campo, texto_livre]):
                                st.warning("Especifique pelo menos um critério de busca.")
                            else:
                                with st.spinner("Buscando..."):
                                    docs = busca_combinada(
                                        vector_store, campo, valor_campo, texto_livre, num_results
                                    )
                                    st.session_state.documentos_contexto = docs # armazena os documentos encontrados DA BUSCA no estado da sessão
                                    if not docs:
                                        st.warning("Nenhum resultado encontrado.")
                    with col2:
                        if st.button("🗑️ Limpar Tudo", key="limpar_tudo"):
                            limpar_tudo()

                # Documentos no sidebar
                if st.session_state.documentos_contexto:
                    with st.expander(f"📚  Total de Documentos Encontrados : {len(st.session_state.documentos_contexto)}", expanded=False):
                        st.markdown(f"Total: {len(st.session_state.documentos_contexto)}")
                        doc_titles = [f"Assunto : {doc.metadata.get('assunto', 'Sem assunto')}" 
                                    for i, doc in enumerate(st.session_state.documentos_contexto)]
                        selected_doc = st.selectbox("Selecione um documento:", doc_titles)
                        
                        if selected_doc:
                            doc_index = doc_titles.index(selected_doc)
                            doc = st.session_state.documentos_contexto[doc_index]
                            st.text_area(
                                "Conteúdo:",
                                doc.page_content,
                                height=300,
                                key=f"text_{doc_index}"
                            )
            with tab2:
                  with st.expander(label="🔗 Arraste seus arquivos aqui", expanded=True):
                    uploaded_files = st.file_uploader(
                        label="",
                        type=["pdf", "docx", "txt","csv"],
                        accept_multiple_files=True,
                        help="Suporta múltiplos arquivos PDF, DOCX,TXT,csv",
                    )
                    
                    # Sincronizar com o estado da sessão
                    if uploaded_files:
                        # Obter nomes dos arquivos atualmente no uploader
                        current_file_names = [f.name for f in uploaded_files]
                        
                        # Remover do estado arquivos que não estão mais no uploader
                        st.session_state.uploaded_files = [
                            f for f in st.session_state.uploaded_files 
                            if f['name'] in current_file_names
                        ]
                        
                        # Adicionar novos arquivos
                        for uploaded_file in uploaded_files:
                            if not any(f['name'] == uploaded_file.name for f in st.session_state.uploaded_files):
                                try:
                                    processed_file = process_uploaded_file(uploaded_file)
                                    st.session_state.uploaded_files.append(processed_file)                                
                                except Exception as e:
                                    st.error(f"Erro ao processar arquivo {uploaded_file.name}: {str(e)}")
                    else:
                        # Se não há arquivos no uploader, limpar o estado também
                        st.session_state.uploaded_files = []

                    # No seu código Streamlit, adicione na sidebar:
                    if st.button("🔄 Converter banco de dados SQLite para vectorstore FAISS"):
                        vector_store = converter_sqlite_para_vectorstore()
                        if vector_store:
                            st.session_state.vector_store = vector_store
                            st.rerun()
               
                    on_sqlite_troggle= st.toggle("Activate feature",
                        key="on_sqlite_troggle_key",
                        help="Ativar ou desativar a conversão do banco de dados SQLite para vectorstore FAISS",
                        value= st.session_state.on_sqlite_troggle)
                    if on_sqlite_troggle:
                        st.session_state.on_sqlite_troggle = True
                    else:   
                        st.session_state.on_sqlite_troggle = False

            with tab3:
                with st.expander("⚙️ Configurações", expanded=True):
                
                    model = st.selectbox(
                        "Modelo de IA:",
                        ["gemini-2.0-flash-thinking-exp-01-21"] + models, 
                        help="Selecione o tipo de modelo de llm",
                        index=0,  
                        key="model_key"                                              
                    )

                    # Atualizar o estado quando o valor mudar
                    st.session_state.model_selecionado = model

                    st.divider()
        
                    # Configurações de Busca
                    st.markdown("**🔍 Configurações de Busca**")
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        search_type = st.selectbox(
                            "Tipo de Busca:",
                            ["similarity", "mmr", "similarity_score_threshold"],
                            index=["similarity", "mmr", "similarity_score_threshold"].index(
                                st.session_state.search_config['search_type']
                            ),
                            help="Tipo de algoritmo de busca",
                            key="search_type_key"
                        )
                        st.session_state.search_config['search_type'] = search_type
                        
                        k = st.number_input(
                            "K (documentos retornados):",
                            min_value=1,
                            max_value=50,
                            value=st.session_state.search_config['k'],
                            help="Número de documentos a retornar",
                            key="k_key"
                        )
                        st.session_state.search_config['k'] = k
                        
                        fetch_k = st.number_input(
                            "Fetch K:",
                            min_value=1,
                            max_value=100,
                            value=st.session_state.search_config['fetch_k'],
                            help="Número de documentos para buscar antes da filtragem",
                            key="fetch_k_key"
                        )
                        st.session_state.search_config['fetch_k'] = fetch_k
                    
                    with col2:
                        lambda_mult = st.slider(
                            "Lambda Mult:",
                            min_value=0.0,
                            max_value=1.0,
                            value=st.session_state.search_config['lambda_mult'],
                            step=0.1,
                            help="Controla diversidade vs relevância (apenas para MMR)",
                            key="lambda_mult_key"
                        )
                        st.session_state.search_config['lambda_mult'] = lambda_mult
                        
                        score_threshold = st.slider(
                            "Score Threshold:",
                            min_value=0.0,
                            max_value=1.0,
                            value=st.session_state.search_config['score_threshold'],
                            step=0.1,
                            help="Limiar mínimo de pontuação",
                            key="score_threshold_key"
                        )
                        st.session_state.search_config['score_threshold'] = score_threshold
                    
                    st.divider()
                    
                    # Configurações de Geração
                    st.markdown("**🤖 Configurações de Geração**")
                    
                    col3, col4 = st.columns(2)
                    
                    with col3:
                        temperature = st.slider(
                            "Temperature:",
                            min_value=0.0,
                            max_value=2.0,
                            value=st.session_state.generation_config['temperature'],
                            step=0.1,
                            help="Controla a criatividade da resposta",
                            key="temperature_key"
                        )
                        st.session_state.generation_config['temperature'] = temperature
                        
                        top_p = st.slider(
                            "Top P:",
                            min_value=0.0,
                            max_value=1.0,
                            value=st.session_state.generation_config['top_p'],
                            step=0.1,
                            help="Nucleus sampling",
                            key="top_p_key"
                        )
                        st.session_state.generation_config['top_p'] = top_p
                        
                        top_k = st.number_input(
                            "Top K:",
                            min_value=1,
                            max_value=100,
                            value=st.session_state.generation_config['top_k'],
                            help="Top-k sampling",
                            key="top_k_key"
                        )
                        st.session_state.generation_config['top_k'] = top_k
                    
                    with col4:
                        max_output_tokens = st.number_input(
                            "Max Output Tokens:",
                            min_value=100,
                            max_value= 8192,
                            value=st.session_state.generation_config['max_output_tokens'],
                            step=100,
                            help="Máximo de tokens na resposta",
                            key="max_tokens_key"
                        )
                        st.session_state.generation_config['max_output_tokens'] = max_output_tokens
                        
                        candidate_count = st.number_input(
                            "Candidate Count:",
                            min_value=1,
                            max_value=5,
                            value=st.session_state.generation_config['candidate_count'],
                            help="Número de candidatos de resposta",
                            key="candidate_count_key"
                        )
                        st.session_state.generation_config['candidate_count'] = candidate_count
                        
                        stop_sequences = st.text_area(
                            "Stop Sequences:",
                            value="\n".join(st.session_state.generation_config['stop_sequences']),
                            help="Sequências que param a geração (uma por linha)",
                            key="stop_sequences_key"
                        )
                        # Converter texto em lista
                        st.session_state.generation_config['stop_sequences'] = [
                            seq.strip() for seq in stop_sequences.split('\n') if seq.strip()
                        ]
                    
                    st.divider()
                    
                    # Botões de controle
                    col5, col6, col7 = st.columns(3)
                    
                    with col5:
                        if st.button("🔄 Resetar Configurações", key="reset_config"):
                            # Resetar para valores padrão
                            st.session_state.search_config = {
                                'search_type': 'mmr',
                                'k': 10,
                                'fetch_k': 20,
                                'lambda_mult': 0.5,
                                'score_threshold': 0.4
                            }
                            st.session_state.generation_config = {
                                'temperature': 0.1,
                                'top_p': 0.9,
                                'top_k': 40,
                                'max_output_tokens': 8192,
                                'candidate_count': 1,
                                'stop_sequences': []
                            }
                            st.rerun()
                    
                    with col6:
                        if st.button("💾 Salvar Config", key="save_config"):
                            # Salvar configurações em arquivo (opcional)
                            config_data = {
                                'search_config': st.session_state.search_config,
                                'generation_config': st.session_state.generation_config
                            }
                            with open('user_config.json', 'w') as f:
                                json.dump(config_data, f, indent=2)
                            st.success("Configurações salvas!")
                    
                    with col7:
                        if st.button("📁 Carregar Config", key="load_config"):
                            # Carregar configurações de arquivo (opcional)
                            try:
                                with open('user_config.json', 'r') as f:
                                    config_data = json.load(f)
                                st.session_state.search_config = config_data['search_config']
                                st.session_state.generation_config = config_data['generation_config']
                                st.success("Configurações carregadas!")
                                st.rerun()
                            except FileNotFoundError:
                                st.warning("Arquivo de configuração não encontrado!")
                                       
        with colChatbot:
            # # Área principal do chat
            # st.markdown("""<center><h2> 💬 Chat Assistente Jurídico</h2><center>""", unsafe_allow_html= True)

            containerChatbot= st.container(height=700, border=True)
            
            with containerChatbot:

                # Exibir mensagens do chat
                for message in st.session_state.messages:
                    with st.chat_message("user" if message["role"] == "user" else "assistant", 
                                        avatar="👨" if message["role"] == "user" else "⚖️"):                     
                    
                        st.markdown(f"""{message["content"]}""", unsafe_allow_html=True)
                
                            # Chat input
            if prompt := st.chat_input("O que você gostaria de perguntar?"):
                st.session_state.messages.append({"role": "user", "content": prompt})                
                
                with containerChatbot:
                    with st.chat_message("user", avatar="👨"):
                        st.markdown(prompt)  
                    
                    with st.spinner("🤔 Pensando..."):                        
                        if campo:
                            filterName = campo
                            filterValue = valor_campo
                        else:
                            filterName = None
                            filterValue = None  
                    
                        prompt_formatado =  gerar_resposta(
                                                    prompt, 
                                                    vector_store,
                                                    st.session_state.documentos_contexto,
                                                    st.session_state.chat_history, 
                                                    st.session_state.uploaded_files 
                                                    , filterName,
                                                    filterValue,
                                                    st.session_state.model_selecionado                                                 
                                                    )         
                        #resposta = st.session_state.chat_session.send_message(prompt_formatado) #, history=st.session_state.chat_history  
                        resposta = prompt_formatado # Resposta formatada
                    with st.chat_message("assistant", avatar="⚖️"):
                        
                        #st.text(f"Resposta \n\n:{resposta.text}") 
                        #st.write_stream(stream_data(resposta.text))
                        #st.markdown(f"Resposta \n\n:{resposta.text}", unsafe_allow_html=True)
                        st.markdown(f"""{resposta}""", unsafe_allow_html=True)

                        st.session_state.messages.append({"role": "assistant", "content": f" {resposta}"})
                        st.session_state.chat_history.extend([
                            {'role': 'user', 'content': prompt},
                            {'role': 'assistant', 'content': resposta}
                        ])    
            
                
                
    else:
        st.error("Não foi possível carregar o vector store")


if __name__ == "__main__":
    main()






